# -*- coding: utf-8 -*-
"""
Created on Sun Jan 26 13:38:28 2020

@author: sse
"""
import numpy as np
import pandas as pd
# from pandas.compat import BytesIO
from io import BytesIO
import matplotlib.pyplot as plt 
import csv
from collections import namedtuple
import joblib

import logging
logging.basicConfig(format='%(levelname)s: %(module)s.%(funcName)s(): %(message)s')
font = {'family': 'serif', 'color':  'darkred', 'weight': 'normal', 'size': 14}

import datetime
import os

import seaborn as sns

cwd = os.getcwd()


def pickle_out(model_):
    Now=datetime.datetime.now()
    joblib_fh=os.path.join(cwd,"RandomForesr_Trianed_%s.pickle"%(Now.strftime("%Y%m%d_%Hh%M")))
    joblib.dump(model_,  joblib_fh, compress=3)
    return joblib_fh
    
def pickle_in(joblib_fh):
    joblib_model = joblib.load(joblib_fh)
    return joblib_model  



def station_data(path):
    '''
    read rainfall and flow station data from a file in the path and returns a 
    namedtuple of stations
    '''
    
    fields = ('id', 
          'name', 
          'x', 
          'y')
    stations=[]
    data = namedtuple('data', fields)

    with open(path,"r") as f:
         r = csv.reader(f, delimiter=',')

         headers = next(f)
         for row in r:
               stations.append(data(
               id          = row[0].strip(),
               name        = row[1].strip(),
               x           = float(row[2].strip()),
               y           = float(row[3].strip()))
               )
    return stations          

import math  
def calculateDistance(x1,y1,x2,y2):  
    '''
    calculates the eculidian distance etween to points
    '''
    dist = math.sqrt((x2 - x1)**2 + (y2 - y1)**2)  
    return dist 


def closest_stations(Rainguages, SWguages, sw,n):
    '''
    returns a list containing the flow station and n number of closest 
    rainguages to a flow station sw Rainguages and SWguages are lists of 
    namedtuples of rain gauges and slowguages
    '''  
       
    for st in SWguages:
       #print(st.id, sw)
        if(st.id == sw):
            distance = []
            for rg in Rainguages:
                dist = calculateDistance(st.x, st.y, rg.x, rg.y)
                a = [rg.id, dist]
                distance.append(a)
                a = sorted(distance, key=lambda x: int(x[1]))
                #lst2 = [item[0] for item in a[:5]]
                lst2 = [item[0] for item in a]
            
            lst3 = [st.id]+lst2[:n]
            break
        else:
            lst3 = [st.id]
    
    return lst3     

def set_col_widths(table):
    '''
    set column width in a table
    '''
    from docx import Document
    from docx.shared import Inches
    
    widths = (Inches(0.5), Inches(1), Inches(1), Inches(1))
    for row in table.rows:
        for cell, width in zip(row.cells, widths): 
            cell.width = width
            
def nashsutcliffe(evaluation, simulation):
    """
    Nash-Sutcliffe model efficinecy
        .. math::
         NSE = 1-\\frac{\\sum_{i=1}^{N}(e_{i}-s_{i})^2}{\\sum_{i=1}^{N}(e_{i}-\\bar{e})^2} 
    :evaluation: Observed data to compared with simulation data.
    :type: list
    :simulation: simulation data to compared with evaluation data
    :type: list
    :return: Nash-Sutcliff model efficiency
    :rtype: float
    """
    if len(evaluation) == len(simulation):
        s, e = np.array(simulation), np.array(evaluation)
        mean_observed = np.nanmean(e)
        numerator = np.nansum((e - s) ** 2)
        denominator = np.nansum((e - mean_observed)**2)
        return 1 - (numerator / denominator)

    else:
        logging.warning("evaluation and simulation lists does not have the same length.")
        return np.nan

def correlationcoefficient(evaluation, simulation):
    """
    Correlation Coefficient
        .. math::
         r = \\frac{\\sum ^n _{i=1}(e_i - \\bar{e})(s_i - \\bar{s})}{\\sqrt{\\sum ^n _{i=1}(e_i - \\bar{e})^2} \\sqrt{\\sum ^n _{i=1}(s_i - \\bar{s})^2}}
    :evaluation: Observed data to compared with simulation data.
    :type: list
    :simulation: simulation data to compared with evaluation data
    :type: list
    :return: Corelation Coefficient
    :rtype: float
    """
    if len(evaluation) == len(simulation):
        correlation_coefficient = np.corrcoef(evaluation, simulation)[0, 1]
        return correlation_coefficient
    else:
        logging.warning("evaluation and simulation lists does not have the same length.")
        return np.nan


def rsquared(evaluation, simulation):
    """
    Coefficient of Determination
        .. math::
         r^2=(\\frac{\\sum ^n _{i=1}(e_i - \\bar{e})(s_i - \\bar{s})}{\\sqrt{\\sum ^n _{i=1}(e_i - \\bar{e})^2} \\sqrt{\\sum ^n _{i=1}(s_i - \\bar{s})^2}})^2
    :evaluation: Observed data to compared with simulation data.
    :type: list
    :simulation: simulation data to compared with evaluation data
    :type: list
    :return: Coefficient of Determination
    :rtype: float
    """
 
    
    
    if len(evaluation) == len(simulation):
        s, e = np.array(simulation), np.array(evaluation)
        correlation_matrix = np.corrcoef(e, s)
        correlation_xy = correlation_matrix[0,1]
        return  correlation_xy**2

        
        # p1 = len(s)*np.sum(e*s)-(np.sum(e)*np.sum(s))
        # p2 = len(s)*np.sum(e**2)-(np.sum(e)**2)
        # p3 = len(s)*np.sum(s**2)-(np.sum(s)**2)
        # return p1/(np.sqrt(p2*p3))
        
        # mean_observed = np.nanmean(e)
        # mean_simulated = np.nanmean(s)
        # num = np.nansum((e-mean_observed)*(s-mean_simulated))
        # denm = np.sqrt(np.nansum((e-mean_observed)**2))*np.sqrt(np.nansum((s-mean_simulated)**2))
        # return (num/denm)**2

    else:
        logging.warning("evaluation and simulation lists does not have the same length.")
        return np.nan

def mse(evaluation, simulation):
    """
    Mean Squared Error
        .. math::
         MSE=\\frac{1}{N}\\sum_{i=1}^{N}(e_{i}-s_{i})^2
    :evaluation: Observed data to compared with simulation data.
    :type: list
    :simulation: simulation data to compared with evaluation data
    :type: list
    :return: Mean Squared Error
    :rtype: float
    """

    if len(evaluation) == len(simulation):
        obs, sim = np.array(evaluation), np.array(simulation)
        mse = np.nanmean((obs - sim)**2)
        return mse
    else:
        logging.warning("evaluation and simulation lists does not have the same length.")
        return np.nan

def rmse(evaluation, simulation):
    """
    Root Mean Squared Error
        .. math::
         RMSE=\\sqrt{\\frac{1}{N}\\sum_{i=1}^{N}(e_{i}-s_{i})^2}
    :evaluation: Observed data to compared with simulation data.
    :type: list
    :simulation: simulation data to compared with evaluation data
    :type: list
    :return: Root Mean Squared Error
    :rtype: float
    """
    if len(evaluation) == len(simulation) > 0:
        return np.sqrt(mse(evaluation, simulation))
    else:
        logging.warning("evaluation and simulation lists do not have the same length.")
        return np.nan   

def evaluation_stat(evaluation, simulation):
    
    if len(evaluation) == len(simulation) > 0:
        return rsquared(evaluation, simulation), \
            rmse(evaluation, simulation), \
            nashsutcliffe(evaluation, simulation) 
    else:
        logging.warning("evaluation and simulation lists do not have the same length.")
        return np.nan   
    return 
	
    
def fit_RF(X_train, X_test, y_train, y_test):
    
    ### Random Forest
#    from sklearn.metrics import mean_squared_error
    from sklearn.ensemble import RandomForestRegressor
    
    # # better training resulst for u05
    # RF = RandomForestRegressor(n_estimators=900, max_features= 2, max_depth = 80, 
    #                           min_samples_split=8, min_samples_leaf=5,
    #                           bootstrap=True, n_jobs=-1, verbose = 1)
    
    ## best forl all
    # RF = RandomForestRegressor(n_estimators=300,  max_depth = 80, 
    #                             min_samples_split=12, min_samples_leaf=5,max_features=3,
    #                             bootstrap=True, n_jobs=-1, verbose = 1)
    # {'bootstrap': True, 'max_depth': 80, 'max_features': 3, 'min_samples_leaf': 5, 
    #  'min_samples_split': 12, 'n_estimators': 300}
    
    ## for U11
    # RF = RandomForestRegressor(n_estimators=1000,  max_depth = 80, 
    #                             min_samples_split=12, min_samples_leaf=5,max_features=3,
    #                             bootstrap=True, n_jobs=-1, verbose = 1)
    
        ## less traing result for U05
    # RF = RandomForestRegressor(n_estimators=200, max_features= 'sqrt', 
    #                           max_depth = 10, 
    #                           min_samples_split=2, min_samples_leaf=10,
    #                           bootstrap=True, n_jobs=-1, verbose = 1)
    
    # ## 2nd best for all
    # RF = RandomForestRegressor(bootstrap=True, criterion='mse', max_depth=90,
    #                   max_features='sqrt', max_leaf_nodes=None,
    #                   min_impurity_decrease=0.0, min_impurity_split=None,
    #                   min_samples_leaf=10, min_samples_split=2,
    #                   min_weight_fraction_leaf=0.0, n_estimators=900,
    #                   n_jobs=-1, oob_score=False, random_state=None,
    #                   verbose=1, warm_start=False)
    
   # ## calibrated for radar raiffall
   #  RF = RandomForestRegressor(bootstrap=True, criterion='mse', max_depth=80,
   #                    max_features=3, max_leaf_nodes=None,
   #                    min_impurity_decrease=0.0, min_impurity_split=None,
   #                    min_samples_leaf=3, min_samples_split=10,
   #                    min_weight_fraction_leaf=0.0, n_estimators=300,
   #                    n_jobs=-1, oob_score=False, random_state=None,
   #                    verbose=1, warm_start=False)

   # ## calibrated for radar raiffall2
    RF = RandomForestRegressor(bootstrap=True, criterion = 'squared_error' max_depth=70,
                       max_features=4, max_leaf_nodes=None,
                       min_impurity_decrease=0.0, 
                       min_samples_leaf=5, min_samples_split=14,
                       min_weight_fraction_leaf=0.0, n_estimators=200,
                       n_jobs=-1, oob_score=False, random_state=None,
                       verbose=1, warm_start=False)

#    
    RF.fit(X_train,y_train)
    
    train_pred = RF.predict(X_train)
    test_pred = RF.predict(X_test)

    return  train_pred,  test_pred, RF
    

def plot_regression_result(method, name, df2plot,r2, color, gaps):
    


        memfile = BytesIO()
        folder = r"D:\VUB\Python codes"
        # stations with level measurment
        WL = ['U06','C02','U09']
        # # reshape the training data
        # obs = obs.reshape(-1,1)
        # est = est.reshape(-1,1)
        
        fig, ax = plt.subplots(figsize=(8,5))
        #fig = plt.figure(figsize=(16,9))
        plt.xlabel("Date", fontdict=font)
        if name in WL:
            plt.ylabel("Water Level [mm]", fontdict=font)  
        else:
            plt.ylabel("Flow [l/s]", fontdict=font) 
        
        if gaps:
            # df2plot = pd.DataFrame(data = np.hstack((obs, est)), index=index,
            #                        columns=['obs', 'est'])
            index = df2plot.index
            index5T = pd.date_range(freq='05T', start=index[0], end=index[-1])
            df2plot = df2plot.reindex(index=index5T)
            df2plot.dropna()
           
            import matplotlib.dates as mdates
            ax.plot( df2plot['observed'], color=color[0],  lw=2.0, label = 'observed')
            ax.plot( df2plot['estimated'], color=color[1], lw=1.5, label = 'estimated')
            ax.xaxis.set_major_locator(mdates.MonthLocator(interval=2))
            ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d'))
                   
        else:
            #ch3ck if obs is a dataframe or not
            # if isinstance(obs, pd.DataFrame):
            #     obs = obs.values
            
            obs = df2plot.observed.values
            est = df2plot.estimated.values
            plt.plot(obs, color=color[0],  lw=2.5, label = 'observed')
            plt.plot(est, color=color[1], lw=1.5, label = 'estimated')
            
            
        plt.text(0.2, 0.9,'{}\u00b2 = {:.2f}'.format('R',r2), ha='center', 
                 va='center', transform=ax.transAxes, fontdict=font)
        name_fig = name+'_'+method    
        
        # ## To plot with date instead of numbers when avoiding plotting with gaps
        # fig.canvas.draw()
        # lables = [item.get_text() for item in ax.get_xticklabels()]
        # del lables[0]
        # del lables[-1]
        # print('lables :', lables)
        # dates = df2plot.dropna().index.to_list()
        # print(len(dates))
        # list_numbers = [ int(x) for x in lables ]
        # l = [dates[i] for i in list_numbers]
        # print('list :',l)
        # l = [0]+l
        # ax.set_xticklabels(l);
        # fig.autofmt_xdate()
        
        plt.title(name_fig, fontdict=font)
        plt.tick_params(axis='both', which='major', labelsize=12)
        plt.tick_params(axis='x', rotation=45)
        plt.legend(loc="best")
        plt.show()
        fig.tight_layout()
        fig.savefig(memfile)
        # fig.savefig(os.path.join(r'd:\VUB\Modelling_result\plots', name_fig+'.png'), dpi = 600)
        return memfile

#============================================================================
def wirte_results_summary_to_word2(name,lag,input_len_train, input_features_len,
                                   input_list,input_len_test, plots_list,result,
                                   models,rolling_sum_window, folder):
    
    
    from docx import Document
    from docx.shared import Inches, Pt
### Write the results in to a word document
    doc = Document()
    doc.add_heading('Random Forest Model Results', 0)
    
    p = doc.add_paragraph()
    style='ListBullet'
    p.add_run('Summary of regression result ').bold  = True
    doc.add_paragraph('Station: '+name, style=style)
    doc.add_paragraph('Lag time in minutes: '+str(lag), style=style)
    doc.add_paragraph('Lag for rainfall rolling sum in minutes: '+str(rolling_sum_window*5), style=style)
    
    doc.add_paragraph('Training data length: '+str(input_len_train), style=style)
    doc.add_paragraph('Test data length: '+str(input_len_test), style=style)
    doc.add_paragraph('Number of input features: '+str(input_features_len), style=style)
    doc.add_paragraph('List of futures: '+str(input_list), style=style)
    
    

    #print results table
    t = doc.add_table(result.shape[0]+1, result.shape[1])
    t.style = 'Colorful List Accent 5'
    #add the header rows
    for j in range(result.shape[-1]):
        t.cell(0,j).text = result.columns[j]
    # add the rest of the Dataframe
    result = result.values
    for i in range(result.shape[0]):
        t.cell(i+1,0).text =  str(result[i,0])
        for j in range(result.shape[-1]-1):
            t.cell(i+1,j+1).text =  str.format('{0:.2f}', result[i,j+1])

    ##change font sstyle and size for the table
    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(10)
                    font.name = 'Calibri Light'
            
    doc.add_page_break()
    

    for plot in plots_list:
        doc.add_picture(plot, width=Inches(5.5))
        plot.close()
    doc.add_page_break()
    doc.add_heading('Model details', level=2)
    doc.add_paragraph(str(models))
    
    namelist = [name,'lag{0}min'.format(lag), 'RFrollingsum{0}timesteps'.format(rolling_sum_window)]
    separator = '_'
    fname = separator.join(namelist)
    

    doc.save(os.path.join(folder,'{0}.docx'.format(fname)))
    
def wirte_all_stn_results(result, folder,rolling_sum_window):
    
    from docx import Document
    from docx.shared import Inches,Pt
### Write the results in to a word document
    doc = Document()
    doc.add_heading('Random Forest results for all statipons', 0)
    p = doc.add_paragraph()
    p.add_run('Random Forest Model Results for training and test sets: ').bold  = True
    style='ListBullet'
    
    s=doc.add_paragraph('Coefficient of Determination (R squared),',style=style)
    s.paragraph_format.left_indent = Inches(0.5)
    s=doc.add_paragraph('Root Mean Squared Error (RMSE) and', style=style)
    s.paragraph_format.left_indent = Inches(0.5)
    s=doc.add_paragraph('Nash-Sutcliffe Efficiency Coefficient (NSE) ', style=style)
    s.paragraph_format.left_indent = Inches(0.5)

    #print results table
    t = doc.add_table(result.shape[0]+1, result.shape[1])
    t.style = 'Colorful List Accent 5'
    t.Title = 'Results of all stations'
    #add the header rows
    for j in range(result.shape[-1]):
        t.cell(0,j).text = result.columns[j]
    # add the rest of the Dataframe
    result = result.values
    for i in range(result.shape[0]):
        t.cell(i+1,0).text =  str(result[i,0])
        for j in range(result.shape[-1]-1):
            t.cell(i+1,j+1).text =  str.format('{0:.2f}', result[i,j+1])

    ##change font sstyle and size for the table
    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(12)
                    font.name = 'Calibri Light'
            

    # doc.save(os.path.join(folder,'{0}_RFrollingsum{1}.docx'.format("All_stn_result",rolling_sum_window)))
    doc.save(os.path.join(folder,'{0}_RFrollingsum.docx'.format(rolling_sum_window)))


def standardize_fit(y):
    y_standard = (y-np.mean(y))/np.std(y)
    return y_standard, np.mean(y), np.std(y)

def standardize(y, mean, std):
    y_standard = (y-mean)/std
    return y_standard

def normalize_fit(y):
    y_normal = (y-np.min(y))/(np.max(y)-np.min(y))
    return y_normal, np.max(y), np.min(y)

def normalize(y,ymax,ymin):
    y_normal = (y-ymin)/(ymax-ymin)
    return y_normal


def HighFlowSeries(obs, est, threshold):
    '''
    identify high flows ffrom a numpy arry (observation) which are greater
    or equal to the value of a cummulative treshholl and correspoinding 
    array from the estimate array
    '''
    x = np.sort(obs)
    x_th = x[int(threshold*x.size)]
    b = np.where(obs>=x_th)[0]
    x_highflow = obs[b]
    
    y_highflow = est[b]
    
    return x_highflow, y_highflow

##
# -*- coding: utf-8 -*-
"""
Created on Sat May  9 20:14:13 2020

@author: sse
"""


def thresholding_algo(y, lag, threshold, influence):
    '''
    Smoothed z-score algorithm (peak detection with robust threshold)
    from http://stackoverflow.com/a/22640362/6029703
    
    lag = the lag of the moving window, 
    threshold = the z-score at which the algorithm signals and 
    influence = the influence (between 0 and 1) of new signals on the mean and standard deviation
    '''
    
    signals = np.zeros(len(y))
    filteredY = np.array(y)
    avgFilter = [0]*len(y)
    stdFilter = [0]*len(y)
    if lag > len(y):
        lag = len(y)
    avgFilter[lag - 1] = np.mean(y[0:lag])
    stdFilter[lag - 1] = np.std(y[0:lag])
    
    for i in range(lag, len(y)):
        if abs(y[i] - avgFilter[i-1]) > threshold * stdFilter [i-1]:
            if y[i] > avgFilter[i-1]:
                signals[i] = 1
            else:
                signals[i] = -1

            filteredY[i] = influence * y[i] + (1 - influence) * filteredY[i-1]
            avgFilter[i] = np.mean(filteredY[(i-lag):i])
            stdFilter[i] = np.std(filteredY[(i-lag):i])
        else:
            signals[i] = 0
            filteredY[i] = y[i]
            avgFilter[i] = np.mean(filteredY[(i-lag):i])
            stdFilter[i] = np.std(filteredY[(i-lag):i])

    return dict(signals = np.asarray(signals),
                avgFilter = np.asarray(avgFilter),
                stdFilter = np.asarray(stdFilter))



def Identify_peaks(dt):
    '''
    identify peaks from a series which are greater than a given threshols value
    and return a df with only peak values (values greater than the threshold value
    '''
    
   #  data = dt.values.copy()
   #  x = np.sort(data) 
   #  y1 = 0.9
   #  x2 = x[int(y1*x.size)]
   
   # # replace all values less than x2 by x2
   #  #mask = x < x2
   #  #x[mask] = x2
    
   #  mask = dt < x2
   #  # column_name = 'flow'
   #  dt.loc[mask] = x2      
            
            
    index = dt.index.values
    y = np.array(dt, dtype=pd.Series) # np array from the Series
    
    # Split the array based on peaks
    lag = 30
    threshold = 5
    influence = 0

    # Run algo with settings from above
    result = thresholding_algo(y, lag=lag, threshold=threshold, influence=influence)
    
    b = np.where(result['signals'] == 0)[0]   
    # replace the the value less than the treshold by nan               
    y[b] = np.nan
    y_list = map(lambda x: x, y)
    
    # create a series from the numpy array
    df = pd.Series(y_list, index=index, dtype=np.float64)
    
    # Convert to sparse then query index to find block locations
    sparse_ts = df.to_sparse()
    block_locs = zip(sparse_ts.sp_index.blocs, sparse_ts.sp_index.blengths)

    # Map the sparse blocks back to the dense timeseries
    blocks = [df.iloc[start:(start + length - 1)] for (start, length) in block_locs]
    k = 0
    df_peak = pd.DataFrame()
    for block in blocks:
        # block.name = 'flow'
        df2 = pd.DataFrame(block)
        df2.index.names = ['date']
        df_peak = pd.concat([df_peak, df2], axis=0)
        k = k+1
    return df_peak


def make_confusion_matrix(cf,
                          group_names=None,
                          categories='auto',
                          count=True,
                          percent=True,
                          cbar=True,
                          xyticks=True,
                          xyplotlabels=True,
                          sum_stats=True,
                          figsize=None,
                          cmap='Blues',
                          title=None):

    '''
    This function will make a pretty plot of an sklearn Confusion Matrix cm using a Seaborn heatmap visualization.
    Arguments
    ---------
    cf:            confusion matrix to be passed in
    group_names:   List of strings that represent the labels row by row to be shown in each square.
    categories:    List of strings containing the categories to be displayed on the x,y axis. Default is 'auto'
    count:         If True, show the raw number in the confusion matrix. Default is True.
    normalize:     If True, show the proportions for each category. Default is True.
    cbar:          If True, show the color bar. The cbar values are based off the values in the confusion matrix.
                   Default is True.
    xyticks:       If True, show x and y ticks. Default is True.
    xyplotlabels:  If True, show 'True Label' and 'Predicted Label' on the figure. Default is True.
    sum_stats:     If True, display summary statistics below the figure. Default is True.
    figsize:       Tuple representing the figure size. Default will be the matplotlib rcParams value.
    cmap:          Colormap of the values displayed from matplotlib.pyplot.cm. Default is 'Blues'
                   See http://matplotlib.org/examples/color/colormaps_reference.html
                   
    title:         Title for the heatmap. Default is None.
    '''


    # CODE TO GENERATE TEXT INSIDE EACH SQUARE
    blanks = ['' for i in range(cf.size)]

    if group_names and len(group_names)==cf.size:
        group_labels = ["{}\n".format(value) for value in group_names]
    else:
        group_labels = blanks

    if count:
        group_counts = ["{0:0.0f}\n".format(value) for value in cf.flatten()]
    else:
        group_counts = blanks

    if percent:
        group_percentages = ["{0:.2%}".format(value) for value in cf.flatten()/np.sum(cf)]
    else:
        group_percentages = blanks

    box_labels = [f"{v1}{v2}{v3}".strip() for v1, v2, v3 in zip(group_labels,group_counts,group_percentages)]
    box_labels = np.asarray(box_labels).reshape(cf.shape[0],cf.shape[1])


    # CODE TO GENERATE SUMMARY STATISTICS & TEXT FOR SUMMARY STATS
    if sum_stats:
        #Accuracy is sum of diagonal divided by total observations
        accuracy  = np.trace(cf) / float(np.sum(cf))

        #if it is a binary confusion matrix, show some more stats
        if len(cf)==2:
            #Metrics for Binary Confusion Matrices
            precision = cf[1,1] / sum(cf[:,1])
            recall    = cf[1,1] / sum(cf[1,:])
            f1_score  = 2*precision*recall / (precision + recall)
            stats_text = "\n\nAccuracy={:0.3f}\nPrecision={:0.3f}\nRecall={:0.3f}\nF1 Score={:0.3f}".format(
                accuracy,precision,recall,f1_score)
        else:
            stats_text = "\n\nAccuracy={:0.3f}".format(accuracy)
    else:
        stats_text = ""


    # SET FIGURE PARAMETERS ACCORDING TO OTHER ARGUMENTS
    if figsize==None:
        #Get default figure size if not set
        figsize = plt.rcParams.get('figure.figsize')

    if xyticks==False:
        #Do not show categories if xyticks is False
        categories=False


    # MAKE THE HEATMAP VISUALIZATION
    plt.figure(figsize=figsize)
    sns.heatmap(cf,annot=box_labels,fmt="",cmap=cmap,cbar=cbar,xticklabels=categories,yticklabels=categories)

    if xyplotlabels:
        plt.ylabel('True label')
        plt.xlabel('Predicted label' + stats_text)
    else:
        plt.xlabel(stats_text)
    
    if title:
        plt.title(title)
